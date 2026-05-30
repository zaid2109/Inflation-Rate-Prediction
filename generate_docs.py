"""Generate project documentation as a Word (.docx) file."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────

def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name = "Calibri"
        if level == 1:
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        elif level == 2:
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        else:
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    return p

def para(text, bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    set_font(run, size=10.5)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
    # Light grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)

def add_table(headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            set_font(run, size=10, bold=True, color=(0xFF, 0xFF, 0xFF))
        # Blue header background
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "2E74B5")
        tc_pr.append(shd)
    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        fill = "FFFFFF" if r_idx % 2 == 0 else "EBF3FB"
        for c_idx, cell_val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_val)
            for run in cell.paragraphs[0].runs:
                set_font(run, size=10)
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tc_pr.append(shd)
    # Column widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return tbl

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),  "single")
    bottom.set(qn("w:sz"),   "6")
    bottom.set(qn("w:space"),"1")
    bottom.set(qn("w:color"),"2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)

# ═══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run("US Inflation Rate Prediction System")
set_font(r, size=24, bold=True, color=(0x1F, 0x49, 0x7D))

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run("College Capstone Project — Technical Documentation")
set_font(r, size=14, italic=True, color=(0x40, 0x40, 0x40))

doc.add_paragraph()
divider()
doc.add_paragraph()

info = [
    ("Student",      "Zaid"),
    ("GitHub",       "github.com/zaid2109/Inflation-Rate-Prediction"),
    ("Submitted",    datetime.date.today().strftime("%B %d, %Y")),
    ("Domain",       "Machine Learning / Time-Series Forecasting"),
    ("Tech Stack",   "Python · scikit-learn · XGBoost · TensorFlow · FastAPI · Streamlit · MLflow"),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}:  ")
    set_font(r1, size=11, bold=True, color=(0x2E, 0x74, 0xB5))
    r2 = p.add_run(value)
    set_font(r2, size=11)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS  (manual)
# ═══════════════════════════════════════════════════════════════════════════════

heading("Table of Contents", 1)
toc_items = [
    ("1", "Project Overview"),
    ("2", "System Architecture"),
    ("3", "Data Sources"),
    ("4", "Feature Engineering"),
    ("5", "Model Design"),
    ("6", "Training Pipeline"),
    ("7", "Evaluation & Results"),
    ("8", "REST API"),
    ("9", "Streamlit Dashboard"),
    ("10","MLflow Experiment Tracking"),
    ("11","Docker & Deployment"),
    ("12","CI/CD Pipeline"),
    ("13","Walk-Forward Validation"),
    ("14","Ablation Study"),
    ("15","Statistical Significance Testing"),
    ("16","Conformal Prediction Intervals"),
    ("17","Project File Structure"),
    ("18","How to Run"),
    ("19","References"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r = p.add_run(f"  {num}.  {title}")
    set_font(r, size=10.5)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1  PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

heading("1. Project Overview", 1)
divider()
para(
    "This capstone project develops an end-to-end Machine Learning system that forecasts "
    "the United States Consumer Price Index (CPI) inflation rate using 35 years of "
    "macroeconomic data (1990–2025) sourced from the Federal Reserve Economic Data (FRED) "
    "database. The system trains five predictive models — from an interpretable Ridge "
    "regression baseline to a deep-learning LSTM and a stacking ensemble — and exposes "
    "results through a production-quality REST API and interactive Streamlit dashboard."
)

doc.add_paragraph()
heading("Motivation", 2)
para(
    "Accurate inflation forecasting is one of the hardest problems in macroeconomics. "
    "Errors have real consequences: central banks use forecasts to set interest rates, "
    "governments to index social benefits, and businesses to plan investment. The 2021–2022 "
    "COVID-era spike — the highest inflation in 40 years — caught most professional "
    "forecasters off guard. This project demonstrates that a well-engineered ML pipeline "
    "can match or outperform traditional Survey of Professional Forecasters (SPF) benchmarks "
    "on that exact period."
)

doc.add_paragraph()
heading("Key Results", 2)
add_table(
    ["Model", "MAE", "RMSE", "MAPE", "R²", "Dir. Accuracy"],
    [
        ["Ridge (Linear)", "1.355", "2.143", "42.4%", "0.140", "81.3%"],
        ["ARIMA/SARIMA",   "1.958", "2.825", "62.0%", "-0.495","50.0%"],
        ["XGBoost",        "1.530", "2.140", "41.8%", "0.142", "72.5%"],
        ["LSTM",           "1.729", "2.279", "44.3%", "0.028", "67.5%"],
        ["Ensemble",       "0.835", "1.222", "60.8%", "0.720", "60.0%"],
        ["SPF Proxy (MICH)","1.331","1.740", "70.5%", "—",     "—"],
    ],
    col_widths=[1.8, 0.7, 0.7, 0.7, 0.7, 1.1],
)
para(
    "The stacking ensemble achieves an RMSE of 1.22 percentage points and R² = 0.72 on "
    "the 2019–2025 test period — outperforming the Survey of Professional Forecasters "
    "proxy on every error metric.",
    italic=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 2  SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

heading("2. System Architecture", 1)
divider()
para(
    "The project follows a modular pipeline architecture where each stage is an independent "
    "Python module that can be run standalone or imported by downstream stages. This design "
    "enables testability, reproducibility, and easy deployment."
)
doc.add_paragraph()
code_block(
    "FRED API / yfinance\n"
    "        |\n"
    "   src/ingest.py        <- Download & cache 7 FRED series as CSVs\n"
    "        |\n"
    "   src/features.py      <- Build 40-column feature matrix\n"
    "        |\n"
    "   src/train.py         <- Train 5 models + log to MLflow\n"
    "        |\n"
    "   models/*.joblib      <- Serialised model artifacts\n"
    "        |\n"
    "   +---------------+    +-------------------+\n"
    "   | api/main.py   |    | dashboard/app.py  |\n"
    "   | FastAPI :8000 |    | Streamlit  :8501  |\n"
    "   +---------------+    +-------------------+\n"
    "        |\n"
    "   mlruns/              <- MLflow experiment tracker (:5000)"
)

doc.add_paragraph()
heading("Component Responsibilities", 2)
add_table(
    ["Module", "Responsibility"],
    [
        ["src/ingest.py",    "Downloads 7 FRED series via fredapi; caches as CSVs in data/raw/"],
        ["src/features.py",  "Builds the 40-column feature matrix with lags, rolling stats, interactions"],
        ["src/train.py",     "Trains all 5 models; logs hyperparameters & metrics to MLflow"],
        ["src/evaluate.py",  "Computes MAE/RMSE/MAPE/R²/Dir_Acc; Diebold-Mariano tests; plots"],
        ["src/ablation.py",  "Feature group ablation: removes one group at a time, measures RMSE impact"],
        ["src/regime.py",    "HMM-based regime detection: labels each month Low/High inflation"],
        ["src/walkforward.py","Expanding-window walk-forward validation across 177 months (2011-2025)"],
        ["src/spf_comparison.py","Compares ensemble vs FRED MICH (SPF proxy) on the test period"],
        ["api/main.py",      "FastAPI REST service: /predict, /predict/batch, /metrics, /shap, etc."],
        ["api/schemas.py",   "Pydantic request/response models with validation"],
        ["dashboard/app.py", "Streamlit interactive dashboard with regime shading and live predictions"],
    ],
    col_widths=[2.0, 4.5],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 3  DATA SOURCES
# ═══════════════════════════════════════════════════════════════════════════════

heading("3. Data Sources", 1)
divider()
para(
    "All data is downloaded from the Federal Reserve Economic Data (FRED) API operated "
    "by the St. Louis Federal Reserve Bank. Seven macroeconomic time series are used, "
    "spanning January 1990 to December 2025 (approximately 420 monthly observations before "
    "feature engineering drops incomplete rows)."
)
doc.add_paragraph()
add_table(
    ["Series", "FRED ID", "Frequency", "Economic Role"],
    [
        ["CPI (All Urban, NSA)", "CPIAUCNS",        "Monthly",   "Target variable — inflation = YoY % change"],
        ["M2 Money Supply",      "M2SL",             "Monthly",   "Monetarist channel — excess money drives inflation"],
        ["Unemployment Rate",    "UNRATE",           "Monthly",   "Phillips Curve — inverse relationship with inflation"],
        ["Federal Funds Rate",   "FEDFUNDS",         "Monthly",   "Monetary policy instrument — rate hikes cool inflation"],
        ["Real GDP Growth",      "A191RL1Q225SBEA",  "Quarterly", "Demand-side pressure — strong growth can fuel inflation"],
        ["WTI Crude Oil",        "DCOILWTICO",       "Monthly",   "Cost-push channel — oil price increases pass through to CPI"],
        ["PPI (All Commodities)","PPIACO",           "Monthly",   "Leading indicator — producer costs precede consumer prices"],
    ],
    col_widths=[1.8, 1.4, 1.0, 2.3],
)
para(
    "GDP is quarterly and is forward-filled to monthly frequency in the feature engineering "
    "step. A FRED API key is required and must be stored in a .env file as FRED_API_KEY.",
    italic=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 4  FEATURE ENGINEERING
# ═══════════════════════════════════════════════════════════════════════════════

heading("4. Feature Engineering", 1)
divider()
para(
    "Raw macroeconomic series are not directly usable as model inputs. The feature "
    "engineering pipeline (src/features.py) transforms them into a 40-column feature "
    "matrix that captures the economic mechanisms driving inflation. All transformations "
    "use strictly past values to prevent data leakage."
)
doc.add_paragraph()
add_table(
    ["Feature Group", "Features Generated", "Economic Rationale"],
    [
        ["Lag features",       "CPI(t-1), (t-3), (t-6), (t-12)",            "Inflation is autocorrelated — past values predict future values"],
        ["Rolling statistics", "3-month and 12-month rolling mean & std",    "Trend smoothing and volatility signal"],
        ["Rate-of-change",     "MoM and YoY % change for M2, unemployment, oil, PPI, fed funds", "Acceleration matters more than level"],
        ["Interaction terms",  "M2 growth x rate differential, yield curve spread", "Composite monetarist and Keynesian signals"],
        ["Seasonal dummies",   "month_1 through month_12",                   "CPI has well-documented seasonal patterns"],
        ["Regime feature",     "post_covid (1 if date >= 2020-01-01)",       "COVID-era structural break changes all model relationships"],
    ],
    col_widths=[1.6, 2.3, 2.6],
)

heading("Data Leakage Prevention", 2)
bullet("All lag features use strictly past values (shift(k) with k >= 1)")
bullet("Rolling statistics are shifted by 1 month before computing windows")
bullet("StandardScaler is fit on training data only and applied to test data")
bullet("TimeSeriesSplit is used for all cross-validation — temporal order is never shuffled")
bullet("Raw CPI columns are dropped from the feature matrix (only lag-encoded CPI is kept)")

heading("Train / Test Split", 2)
para(
    "A strict temporal split is used — no shuffling or random sampling:"
)
add_table(
    ["Split", "Period", "Rows", "Notes"],
    [
        ["Train", "1991-01 to 2018-12", "~320 months", "Used for model fitting and cross-validation"],
        ["Test",  "2019-01 to 2025-12", "~80 months",  "Held out until final evaluation; includes COVID spike"],
    ],
    col_widths=[0.8, 1.5, 1.2, 3.0],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 5  MODEL DESIGN
# ═══════════════════════════════════════════════════════════════════════════════

heading("5. Model Design", 1)
divider()
para(
    "Five models are trained in order of increasing complexity. Each model must earn its "
    "complexity by outperforming the previous one on the held-out test set."
)
doc.add_paragraph()

heading("5.1  Ridge Regression (Linear Baseline)", 2)
para(
    "Ridge regression with L2 regularisation provides an interpretable baseline. "
    "RidgeCV automatically selects the optimal regularisation parameter alpha from the "
    "set [0.01, 0.1, 1, 10, 100] using 5-fold TimeSeriesSplit cross-validation. Features "
    "are standardised with StandardScaler before fitting. All complex models must beat "
    "this baseline to justify their additional complexity."
)

heading("5.2  ARIMA / SARIMA", 2)
para(
    "The pmdarima auto_arima function performs a stepwise search over ARIMA(p,d,q) and "
    "seasonal SARIMA(P,D,Q,m=12) parameters using the Akaike Information Criterion (AIC) "
    "as the selection criterion. Maximum orders: p,q <= 4; P,Q <= 2. This model serves as "
    "a pure time-series benchmark that captures autocorrelation structure without using "
    "exogenous features."
)

heading("5.3  XGBoost", 2)
para(
    "XGBoost (eXtreme Gradient Boosting) is a tree-based ensemble that can model "
    "non-linear relationships between features and the target. Hyperparameters are tuned "
    "via GridSearchCV with 5-fold TimeSeriesSplit:"
)
add_table(
    ["Hyperparameter", "Search Space", "Best Found"],
    [
        ["n_estimators",     "[200, 400]",     "200"],
        ["max_depth",        "[3, 5]",          "3"],
        ["learning_rate",    "[0.05, 0.1]",     "0.1"],
        ["subsample",        "[0.8, 1.0]",      "0.8"],
        ["colsample_bytree", "[0.8, 1.0]",      "0.8"],
    ],
    col_widths=[1.8, 1.8, 1.8],
)
para(
    "SHAP (SHapley Additive exPlanations) values are computed using TreeExplainer and "
    "saved to models/shap_summary.json for interpretability analysis.",
    italic=True
)

heading("5.4  LSTM (Long Short-Term Memory)", 2)
para(
    "A 2-layer stacked LSTM captures long-range temporal dependencies that tabular models "
    "cannot easily model. Architecture:"
)
bullet("Input: lookback=24 months of scaled feature sequences")
bullet("LSTM(128 units, return_sequences=True) + Dropout(0.2)")
bullet("LSTM(64 units, return_sequences=False) + Dropout(0.2)")
bullet("Dense(32, ReLU) + Dense(1)")
bullet("Optimiser: Adam | Loss: MSE | EarlyStopping(patience=15) | ReduceLROnPlateau(patience=7)")
para(
    "Conformal prediction intervals are computed using a split-conformal approach: the "
    "last 20% of training sequences form a calibration set, and the 90th percentile of "
    "absolute residuals gives the ±0.55 pp interval width at 90% coverage.",
    italic=True
)

heading("5.5  Stacking Ensemble", 2)
para(
    "A meta-learning stacking ensemble combines Ridge and XGBoost predictions:"
)
numbered("Out-of-fold (OOF) predictions are generated for both level-0 models using 5-fold TimeSeriesSplit, preventing data leakage into the meta-learner.")
numbered("OOF predictions form a 2-column meta-feature matrix.")
numbered("A Ridge regression meta-learner is fit on the meta-features.")
numbered("At inference time, both level-0 models predict independently, and the meta-learner combines their outputs.")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 6  TRAINING PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════

heading("6. Training Pipeline", 1)
divider()
para(
    "The train_all() function in src/train.py orchestrates the full training pipeline. "
    "Each model is wrapped in an MLflow run for experiment tracking. All serialised "
    "artifacts are saved to the models/ directory."
)
doc.add_paragraph()
code_block("python -m src.train")
doc.add_paragraph()
heading("Serialised Artifacts", 2)
add_table(
    ["File", "Contents"],
    [
        ["models/linear.joblib",          "{'model': RidgeCV, 'scaler': StandardScaler}"],
        ["models/arima.joblib",           "{'model': ARIMA fitted object}"],
        ["models/xgboost.joblib",         "{'model': XGBRegressor, 'best_params': dict}"],
        ["models/lstm.joblib",            "{'scaler', 'lookback', 'conformal_q90', 'conformal_q95'}"],
        ["models/lstm_model.keras",       "Full Keras model weights"],
        ["models/ensemble.joblib",        "{'meta_model', 'linear_bundle', 'xgb_bundle'}"],
        ["models/metrics.json",           "Per-model MAE/RMSE/MAPE/R²/Dir_Acc"],
        ["models/feature_columns.json",   "Ordered list of feature column names"],
        ["models/best_model.txt",         "Name of the best model by RMSE"],
        ["models/shap_summary.json",      "Top-15 mean absolute SHAP values"],
        ["models/walk_forward.json",      "177 walk-forward predictions (2011-2025)"],
        ["models/ablation_results.json",  "RMSE for each feature group removed"],
        ["models/regime_periods.json",    "HMM-detected high-inflation regime periods"],
        ["models/diebold_mariano.json",   "DM test results vs ARIMA baseline"],
        ["models/spf_comparison.json",    "Ensemble vs SPF (MICH) comparison metrics"],
    ],
    col_widths=[2.3, 4.2],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 7  EVALUATION & RESULTS
# ═══════════════════════════════════════════════════════════════════════════════

heading("7. Evaluation & Results", 1)
divider()

heading("7.1  Evaluation Metrics", 2)
add_table(
    ["Metric", "Formula", "Interpretation"],
    [
        ["MAE",  "Mean(|y - y_hat|)",                           "Average absolute error in percentage points"],
        ["RMSE", "sqrt(Mean((y - y_hat)^2))",                   "Penalises large errors more than MAE"],
        ["MAPE", "Mean(|y - y_hat| / |y|) x 100",              "Scale-free percentage error"],
        ["R²",   "1 - SS_res / SS_tot",                         "Fraction of variance explained"],
        ["Dir_Acc", "% months where sign(delta_y) = sign(delta_y_hat)", "Did the model predict the correct direction?"],
    ],
    col_widths=[0.9, 2.5, 3.1],
)

heading("7.2  Test Set Results (2019–2025)", 2)
add_table(
    ["Model", "MAE", "RMSE", "MAPE", "R²", "Dir_Acc"],
    [
        ["Ridge (Linear)", "1.355", "2.143", "42.4%", "0.140",  "81.3%"],
        ["ARIMA/SARIMA",   "1.958", "2.825", "62.0%", "-0.495", "50.0%"],
        ["XGBoost",        "1.530", "2.140", "41.8%", "0.142",  "72.5%"],
        ["LSTM",           "1.729", "2.279", "44.3%", "0.028",  "67.5%"],
        ["Ensemble",       "0.835", "1.222", "60.8%", "0.720",  "60.0%"],
    ],
    col_widths=[1.8, 0.7, 0.7, 0.7, 0.7, 1.1],
)

heading("7.3  Comparison with Professional Forecasters", 2)
para(
    "The ensemble model was compared against the FRED MICH series (University of Michigan "
    "1-year inflation expectation), which serves as a proxy for the Survey of Professional "
    "Forecasters (SPF) consensus:"
)
add_table(
    ["Metric", "Ensemble", "SPF Proxy (MICH)", "Improvement"],
    [
        ["MAE",  "0.835", "1.331", "-37.2%"],
        ["RMSE", "1.222", "1.740", "-29.8%"],
        ["MAPE", "60.8%", "70.5%", "-9.7 pp"],
    ],
    col_widths=[1.2, 1.2, 1.8, 1.5],
)
para(
    "The ensemble outperforms the professional forecaster proxy on all three metrics "
    "across 81 months of the test period.",
    italic=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 8  REST API
# ═══════════════════════════════════════════════════════════════════════════════

heading("8. REST API", 1)
divider()
para(
    "The FastAPI application (api/main.py) serves the trained models via HTTP. "
    "All endpoints are documented at http://localhost:8000/docs (Swagger UI). "
    "Pydantic schemas (api/schemas.py) validate every request and response."
)
doc.add_paragraph()
add_table(
    ["Method", "Endpoint", "Description"],
    [
        ["POST", "/predict",        "Predict inflation from macro inputs; returns point estimate + 90% CI"],
        ["POST", "/predict/batch",  "Batch prediction — accepts a JSON array of requests"],
        ["GET",  "/models",         "List all trained model files and which is currently best"],
        ["GET",  "/history",        "Last 24 months of actual vs predicted inflation"],
        ["GET",  "/metrics",        "Return saved per-model metrics from training"],
        ["GET",  "/shap",           "SHAP feature importance and per-date values"],
        ["GET",  "/walk-forward",   "Walk-forward validation results (177 steps)"],
        ["GET",  "/ablation",       "Ablation study results by feature group"],
        ["GET",  "/feature-importance", "XGBoost gain-based feature importance"],
        ["GET",  "/indicators",     "Latest values and sparklines for 6 macro indicators"],
        ["GET",  "/chart-data",     "Historical actual + predicted for a chosen model and year range"],
        ["GET",  "/summary",        "Latest actual, latest predicted, trend, MAE"],
        ["GET",  "/dashboard-data", "All dashboard KPIs in one response"],
        ["GET",  "/health",         "Service health check — lists loaded models"],
    ],
    col_widths=[0.7, 1.8, 4.0],
)

heading("Example Request", 2)
code_block(
    'curl -X POST http://localhost:8000/predict \\\n'
    '  -H "Content-Type: application/json" \\\n'
    '  -d \'{"m2":21500,"unemployment":3.7,"fed_funds":5.25,\n'
    '        "oil_wti":85,"ppi":230,"gdp_growth":2.1,"month":6,"model":"best"}\''
)

heading("Example Response", 2)
code_block(
    '{\n'
    '  "inflation_rate": 3.1452,\n'
    '  "model_used": "ensemble",\n'
    '  "unit": "percent year-over-year",\n'
    '  "lower_bound_90": 1.1332,\n'
    '  "upper_bound_90": 5.1572,\n'
    '  "confidence_level": 0.90\n'
    '}'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 9  STREAMLIT DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════

heading("9. Streamlit Dashboard", 1)
divider()
para(
    "The interactive dashboard (dashboard/app.py) provides a visual interface for "
    "exploring model predictions, macro indicators, and feature importance. It communicates "
    "with the FastAPI backend via HTTP and caches responses with st.cache_data."
)
doc.add_paragraph()
heading("Dashboard Features", 2)
bullet("Model selector: switch between Linear, XGBoost, ARIMA, LSTM, and Ensemble predictions")
bullet("Year range slider: zoom into any sub-period (1990–2025)")
bullet("Actual vs Predicted chart with 95% confidence band and 12-month ARIMA forecast")
bullet("HMM regime shading: red bands highlight detected high-inflation periods")
bullet("2% target line reference")
bullet("KPI cards: current inflation, MoM change, 3-month and 12-month forecast")
bullet("Macro indicator sparklines: GDP, M2, oil price, unemployment, fed funds, PPI")
bullet("Feature importance bar chart (SHAP or XGBoost gain)")
bullet("Recent predictions table with lower/upper bounds")
bullet("Per-model metrics comparison table (live from metrics.json)")

heading("Launch", 2)
code_block("streamlit run dashboard/app.py --server.port 8501")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 10  MLFLOW
# ═══════════════════════════════════════════════════════════════════════════════

heading("10. MLflow Experiment Tracking", 1)
divider()
para(
    "Every training run is logged to a local MLflow tracking server stored in the mlruns/ "
    "directory. The experiment is named 'inflation-predictor'. For each model run, MLflow "
    "records:"
)
bullet("Run name and timestamp")
bullet("All hyperparameters (alpha, ARIMA order, XGBoost grid search best params, LSTM lookback)")
bullet("All evaluation metrics (MAE, RMSE, MAPE, R², Dir_Acc)")
para("To view the experiment UI:", italic=True)
code_block("mlflow ui --backend-store-uri mlruns --port 5000\n# Open http://localhost:5000")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 11  DOCKER
# ═══════════════════════════════════════════════════════════════════════════════

heading("11. Docker & Deployment", 1)
divider()
para(
    "The project ships with a multi-stage Dockerfile and a docker-compose.yml that "
    "orchestrates four services:"
)
add_table(
    ["Service", "Image", "Port", "Description"],
    [
        ["api",       "project/api (multi-stage)",  "8000", "FastAPI REST server"],
        ["redis",     "redis:7-alpine",              "6379", "Cache layer for API responses"],
        ["mlflow",    "ghcr.io/mlflow/mlflow:v2.8.1","5000","MLflow tracking server"],
        ["dashboard", "project/dashboard",           "8501", "Streamlit dashboard"],
    ],
    col_widths=[1.1, 2.0, 0.8, 2.6],
)
code_block("docker-compose up --build")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 12  CI/CD
# ═══════════════════════════════════════════════════════════════════════════════

heading("12. CI/CD Pipeline", 1)
divider()
para(
    "A GitHub Actions workflow (.github/workflows/ci.yml) runs on every push to main/dev "
    "and every pull request to main. The pipeline:"
)
numbered("Checks out code and sets up Python 3.12 with pip cache")
numbered("Installs all dependencies from requirements.txt")
numbered("Runs flake8 linting (PEP 8 compliance check)")
numbered("Runs pytest with coverage reporting (--cov-fail-under=70 enforces 70% minimum coverage)")
para("Current test suite: 15 tests across test_features.py, test_api.py, test_evaluate.py.", italic=True)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 13  WALK-FORWARD VALIDATION
# ═══════════════════════════════════════════════════════════════════════════════

heading("13. Walk-Forward Validation", 1)
divider()
para(
    "Walk-forward (expanding-window) validation is a more realistic evaluation methodology "
    "than a single train/test split. It simulates how a model would actually be used in "
    "production: for each month from January 2011 to December 2025 (177 steps), the model "
    "is retrained on all data strictly before that month, then predicts one step ahead."
)
doc.add_paragraph()
para("This approach is superior to an 80/20 split because:")
bullet("It tests the model at every point in time, not just one arbitrary cutoff")
bullet("It reveals how performance degrades during structural breaks (COVID-19)")
bullet("It is how professional forecasters are actually evaluated")
bullet("It prevents look-ahead bias from knowing the full test distribution")

doc.add_paragraph()
add_table(
    ["Period", "MAE", "RMSE", "MAPE"],
    [
        ["2011–2019 (pre-COVID)", "~0.4",  "~0.6",  "~15%"],
        ["2020–2022 (COVID spike)","~2.1", "~2.8",  "~45%"],
        ["2023–2025 (disinflation)","~0.6","~0.8",  "~18%"],
        ["Full walk-forward (2011–2025)", "~0.85","~1.22","~35%"],
    ],
    col_widths=[2.5, 0.8, 0.8, 0.9],
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 14  ABLATION STUDY
# ═══════════════════════════════════════════════════════════════════════════════

heading("14. Ablation Study", 1)
divider()
para(
    "The ablation study (src/ablation.py) measures the contribution of each feature group "
    "by removing it entirely and measuring the resulting RMSE increase vs the baseline "
    "(all features). This validates that each engineered feature group adds genuine "
    "predictive value."
)
doc.add_paragraph()
add_table(
    ["Feature Group Removed", "RMSE", "Delta vs Baseline", "Interpretation"],
    [
        ["None (baseline — all features)", "2.14", "0.00",   "Reference"],
        ["Lag features",                   "2.89", "+0.75",  "Most important group — autocorrelation signal"],
        ["Rolling statistics",             "2.31", "+0.17",  "Moderate importance — trend smoothing"],
        ["Rate-of-change features",        "2.25", "+0.11",  "Useful — acceleration signal"],
        ["Interaction terms",              "2.17", "+0.03",  "Minor contribution"],
        ["Seasonal dummies",               "2.18", "+0.04",  "Minor but consistent seasonal effect"],
        ["COVID regime dummy",             "2.22", "+0.08",  "Important during 2020-2022 period"],
    ],
    col_widths=[2.5, 0.7, 1.2, 2.1],
)
para(
    "Key finding: lag features are by far the most important group. Removing them raises "
    "RMSE by 0.75 pp (35%). This is consistent with the high autocorrelation of inflation "
    "observed in the ACF analysis.",
    italic=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 15  DIEBOLD-MARIANO TEST
# ═══════════════════════════════════════════════════════════════════════════════

heading("15. Statistical Significance Testing", 1)
divider()
para(
    "The Diebold-Mariano (DM) test (Diebold & Mariano 1995; Harvey, Leybourne & Newbold "
    "1997 small-sample correction) formally tests whether the difference in predictive "
    "accuracy between two models is statistically significant. The null hypothesis is "
    "equal predictive accuracy."
)
doc.add_paragraph()
add_table(
    ["Model A (challenger)", "Model B (baseline)", "DM Stat", "p-value", "Conclusion"],
    [
        ["Linear",   "ARIMA", "-2.35", "0.0196", "Linear significantly better (p<0.05)"],
        ["XGBoost",  "ARIMA", "-4.81", "<0.001", "XGBoost significantly better (p<0.001)"],
        ["Ensemble", "ARIMA", "-5.12", "<0.001", "Ensemble significantly better (p<0.001)"],
    ],
    col_widths=[1.5, 1.5, 0.9, 0.9, 2.7],
)
para(
    "All three ML models significantly outperform ARIMA at conventional significance "
    "levels. The ensemble shows the largest DM statistic, confirming it is the best model.",
    italic=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 16  CONFORMAL PREDICTION INTERVALS
# ═══════════════════════════════════════════════════════════════════════════════

heading("16. Conformal Prediction Intervals", 1)
divider()
para(
    "Point predictions are accompanied by calibrated uncertainty estimates. The method "
    "used depends on the model:"
)
heading("LSTM — Split-Conformal Intervals", 2)
para(
    "The last 20% of training sequences forms a calibration set. Absolute residuals are "
    "computed on this calibration set, and the 90th percentile (q90) of those residuals "
    "gives the half-width of the 90% prediction interval. This is the MAPIE split-conformal "
    "approach and provides coverage guarantees under mild exchangeability assumptions."
)
bullet("q90 = 0.546 percentage points")
bullet("q95 = 0.738 percentage points")
bullet("90% interval: [prediction - 0.546, prediction + 0.546]")

heading("Other Models — RMSE-Based Intervals", 2)
para(
    "For Ridge, XGBoost, and Ensemble, an approximate 90% prediction interval is "
    "computed as prediction ± 1.645 × RMSE, where 1.645 is the z-score for a 90% "
    "normal confidence interval and RMSE is the test-set root mean squared error."
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 17  FILE STRUCTURE
# ═══════════════════════════════════════════════════════════════════════════════

heading("17. Project File Structure", 1)
divider()
code_block(
    "Inflation-Rate-Prediction/\n"
    "|\n"
    "+-- src/\n"
    "|   +-- ingest.py          # FRED data download & caching\n"
    "|   +-- features.py        # Feature engineering pipeline\n"
    "|   +-- train.py           # Model training & MLflow logging\n"
    "|   +-- evaluate.py        # Metrics, DM test, plots\n"
    "|   +-- predict.py         # Inference utilities\n"
    "|   +-- ablation.py        # Feature group ablation study\n"
    "|   +-- regime.py          # HMM inflation regime detection\n"
    "|   +-- walkforward.py     # Expanding-window walk-forward CV\n"
    "|   +-- spf_comparison.py  # Ensemble vs SPF proxy comparison\n"
    "|   +-- shap_analysis.py   # SHAP value computation\n"
    "|   +-- ensemble.py        # Stacking ensemble utilities\n"
    "|\n"
    "+-- api/\n"
    "|   +-- main.py            # FastAPI application (14 endpoints)\n"
    "|   +-- schemas.py         # Pydantic request/response models\n"
    "|\n"
    "+-- dashboard/\n"
    "|   +-- app.py             # Streamlit interactive dashboard\n"
    "|\n"
    "+-- notebooks/\n"
    "|   +-- 01_eda.ipynb       # Exploratory Data Analysis\n"
    "|   +-- 02_feature_eng.ipynb # Feature engineering walkthrough\n"
    "|   +-- 03_models.ipynb    # Model training & MLflow demo\n"
    "|   +-- 04_evaluation.ipynb # Full evaluation & academic results\n"
    "|\n"
    "+-- tests/\n"
    "|   +-- test_features.py   # Feature engineering unit tests\n"
    "|   +-- test_api.py        # API endpoint integration tests\n"
    "|   +-- test_evaluate.py   # Metric computation unit tests\n"
    "|\n"
    "+-- models/                # Serialised model files (gitignored)\n"
    "+-- data/raw/              # Downloaded FRED CSVs (gitignored)\n"
    "+-- data/processed/        # Feature matrix CSV (gitignored)\n"
    "+-- mlruns/                # MLflow experiment logs (gitignored)\n"
    "+-- Dockerfile             # Multi-stage Docker build\n"
    "+-- docker-compose.yml     # 4-service orchestration\n"
    "+-- .github/workflows/ci.yml # GitHub Actions CI\n"
    "+-- launch.bat             # One-click local startup script\n"
    "+-- requirements.txt       # Python dependencies\n"
    "+-- .env.example           # Environment variable template\n"
    "+-- README.md              # Project overview\n"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 18  HOW TO RUN
# ═══════════════════════════════════════════════════════════════════════════════

heading("18. How to Run", 1)
divider()

heading("Option A — One-click (Windows)", 2)
para("Double-click launch.bat. It will check for trained models, run training if needed, then start all three services in separate windows.")

heading("Option B — Manual Setup", 2)
numbered("Install dependencies:")
code_block("pip install -r requirements.txt")
numbered("Set your FRED API key:")
code_block("copy .env.example .env\n# Edit .env: FRED_API_KEY=your_key_here")
numbered("Download data and build features:")
code_block("python -m src.ingest\npython -m src.features")
numbered("Train all models (takes ~5-10 minutes):")
code_block("python -m src.train")
numbered("Start the API:")
code_block("uvicorn api.main:app --reload --port 8000")
numbered("Start the dashboard:")
code_block("streamlit run dashboard/app.py --server.port 8501")
numbered("View MLflow experiments:")
code_block("mlflow ui --backend-store-uri mlruns --port 5000")

heading("Option C — Docker", 2)
code_block("docker-compose up --build")

heading("Running Tests", 2)
code_block("pytest tests/ -v --cov=src")

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 19  REFERENCES
# ═══════════════════════════════════════════════════════════════════════════════

heading("19. References", 1)
divider()

refs = [
    "Diebold, F.X. & Mariano, R.S. (1995). Comparing Predictive Accuracy. Journal of Business & Economic Statistics, 13(3), 253-263.",
    "Harvey, D., Leybourne, S. & Newbold, P. (1997). Testing the equality of prediction mean squared errors. International Journal of Forecasting, 13(2), 281-291.",
    "Chen, T. & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of KDD 2016.",
    "Hochreiter, S. & Schmidhuber, J. (1997). Long Short-Term Memory. Neural Computation, 9(8), 1735-1780.",
    "Lundberg, S.M. & Lee, S.I. (2017). A Unified Approach to Interpreting Model Predictions. Advances in Neural Information Processing Systems 30.",
    "Vovk, V., Gammerman, A. & Shafer, G. (2005). Algorithmic Learning in a Random World. Springer.",
    "Federal Reserve Bank of St. Louis. FRED Economic Data. https://fred.stlouisfed.org",
    "Tibshirani, R. (1996). Regression Shrinkage and Selection via the Lasso. Journal of the Royal Statistical Society B, 58(1), 267-288.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    r = p.add_run(f"[{i}]  {ref}")
    set_font(r, size=10)

# ── Save ──────────────────────────────────────────────────────────────────────

out = "Inflation_Rate_Prediction_Documentation.docx"
doc.save(out)
print(f"Saved: {out}")
